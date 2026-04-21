"""
parser.py — Resume File Parser
================================
Concept: Think of this like a universal document translator.
You hand it a PDF or DOCX file, it hands you back clean plain text.

Why do we need this?
- PDFs store text in a binary format — you can't just open and read them like .txt files
- DOCX files are actually ZIP archives containing XML — again, not plain text
- Streamlit gives us raw bytes (not file paths) — so we need in-memory parsing

Key concept — io.BytesIO:
- Normally file parsers expect a file PATH like "/home/user/resume.pdf"
- Streamlit gives us file BYTES (raw binary data in memory)
- io.BytesIO wraps those bytes into a fake "file-like object" that parsers can read
- Analogy: BytesIO is like putting water into a bottle — the water is the bytes,
  the bottle is the file-like wrapper that tools know how to pour from

Common error to watch for:
- PyMuPDF is installed as `PyMuPDF` but imported as `fitz` — this trips everyone up!
  If you get ModuleNotFoundError: No module named 'fitz', run: pip install PyMuPDF
"""

import io
import re
import os
from pathlib import Path


# ── PDF parsing library (installed as PyMuPDF, imported as fitz) ──────────────
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("Warning: PyMuPDF not installed. PDF parsing will be unavailable.")
    print("Fix: pip install PyMuPDF")

# ── DOCX parsing library ───────────────────────────────────────────────────────
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX parsing will be unavailable.")
    print("Fix: pip install python-docx")


class ResumeParser:
    """
    Parses resume files (PDF, DOCX, TXT) into clean plain text.

    Supports two modes:
    1. File path mode  — parse(file_path="resume.pdf")
    2. Bytes mode      — parse(file_content=bytes_data, file_type=".pdf")

    Bytes mode is used by Streamlit since uploaded files arrive as raw bytes.
    """

    # Supported file extensions
    SUPPORTED_FORMATS = {".pdf", ".docx", ".doc", ".txt"}

    def __init__(self):
        # Store the last parsed text so other methods (get_sections, etc.) can use it
        self._last_parsed_text = ""

    # ──────────────────────────────────────────────────────────────────────────
    # PUBLIC API
    # ──────────────────────────────────────────────────────────────────────────

    def parse(self, file_path=None, file_content=None, file_type=None):
        """
        Main entry point. Accepts either a file path OR raw bytes + file type.

        Args:
            file_path    : str or Path — path to resume file on disk
            file_content : bytes — raw file bytes (used with Streamlit uploads)
            file_type    : str — file extension like ".pdf" or ".docx"

        Returns:
            str — cleaned plain text extracted from the resume

        Raises:
            ValueError if neither file_path nor file_content is provided
            ValueError if file format is unsupported
        """
        if file_path is not None:
            # Mode 1: read from disk
            return self._parse_from_path(file_path)
        elif file_content is not None and file_type is not None:
            # Mode 2: parse raw bytes (Streamlit uploads)
            return self._parse_from_bytes(file_content, file_type)
        else:
            raise ValueError(
                "Provide either file_path OR (file_content + file_type)."
            )

    def get_sections(self):
        """
        Detects which standard resume sections are present in the last parsed text.

        Uses regex to look for common section header patterns like:
        "EDUCATION", "Education:", "-- Education --", etc.

        Returns:
            dict — { section_name: bool } indicating presence of each section
        """
        text = self._last_parsed_text
        if not text:
            return {}

        # Each section maps to a list of regex patterns that could represent it
        # re.IGNORECASE makes matching case-insensitive
        section_patterns = {
            "contact":        r"\b(contact|email|phone|address|linkedin|github)\b",
            "summary":        r"\b(summary|objective|profile|about me|overview)\b",
            "experience":     r"\b(experience|employment|work history|career|positions?)\b",
            "education":      r"\b(education|academic|degree|university|college|school)\b",
            "skills":         r"\b(skills?|technical skills?|competenc|expertise|technologies)\b",
            "projects":       r"\b(projects?|portfolio|personal projects?|side projects?)\b",
            "certifications": r"\b(certifications?|certificates?|credentials?|licenses?)\b",
            "achievements":   r"\b(achievements?|accomplishments?|awards?|honors?)\b",
            "publications":   r"\b(publications?|research|papers?|articles?)\b",
        }

        detected = {}
        for section, pattern in section_patterns.items():
            # re.search returns a match object if found, None if not
            detected[section] = bool(re.search(pattern, text, re.IGNORECASE))

        return detected

    def extract_contact_info(self):
        """
        Extracts structured contact information using regex patterns.

        Regex concepts used here:
        - \b  = word boundary (don't match mid-word)
        - \d  = any digit 0-9
        - +   = one or more
        - ?   = zero or one (makes preceding item optional)
        - []  = character class (any char inside brackets)

        Returns:
            dict — { email, phone, linkedin, github } — None if not found
        """
        text = self._last_parsed_text
        if not text:
            return {}

        contact = {}

        # Email: word chars + @ + domain + . + tld
        # Example match: john.doe@gmail.com
        email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        email_match = re.search(email_pattern, text)
        contact["email"] = email_match.group() if email_match else None

        # Phone: handles formats like +91-9876543210, (123) 456-7890, 123.456.7890
        phone_pattern = r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}"
        phone_match = re.search(phone_pattern, text)
        contact["phone"] = phone_match.group().strip() if phone_match else None

        # LinkedIn: matches linkedin.com/in/username style URLs
        linkedin_pattern = r"linkedin\.com/in/[\w\-]+"
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        contact["linkedin"] = linkedin_match.group() if linkedin_match else None

        # GitHub: matches github.com/username style URLs
        github_pattern = r"github\.com/[\w\-]+"
        github_match = re.search(github_pattern, text, re.IGNORECASE)
        contact["github"] = github_match.group() if github_match else None

        return contact

    # ──────────────────────────────────────────────────────────────────────────
    # PRIVATE METHODS — file reading
    # ──────────────────────────────────────────────────────────────────────────

    def _parse_from_path(self, file_path):
        """
        Reads file from disk, then delegates to bytes parser.
        Converting to bytes keeps the parsing logic in one place (_parse_from_bytes).
        """
        file_path = Path(file_path)
        file_type = file_path.suffix.lower()

        if file_type not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported format: {file_type}. "
                f"Supported: {self.SUPPORTED_FORMATS}"
            )

        # Read entire file as bytes — works for both binary (PDF) and text (TXT)
        with open(file_path, "rb") as f:
            content = f.read()

        return self._parse_from_bytes(content, file_type)

    def _parse_from_bytes(self, content, file_type):
        """
        Routes bytes to the correct parser based on file extension.

        Args:
            content   : bytes — raw file content
            file_type : str   — ".pdf", ".docx", ".txt"

        Returns:
            str — cleaned extracted text
        """
        file_type = file_type.lower()

        if file_type == ".pdf":
            raw_text = self._parse_pdf(content)
        elif file_type in (".docx", ".doc"):
            raw_text = self._parse_docx(content)
        elif file_type == ".txt":
            # TXT files: just decode bytes to string
            raw_text = content.decode("utf-8", errors="ignore")
        else:
            raise ValueError(f"Unsupported format: {file_type}")

        # Clean the text and store it for get_sections() / extract_contact_info()
        cleaned = self._clean_text(raw_text)
        self._last_parsed_text = cleaned
        return cleaned

    def _parse_pdf(self, content):
        """
        Extracts text from PDF bytes using PyMuPDF (fitz).

        How PDFs work internally:
        - A PDF is a binary format with pages, fonts, and drawing commands
        - fitz.open() decodes this binary format
        - page.get_text() extracts the text layer from each page
        - Some PDFs are "image-only" scans — those have no text layer at all

        Args:
            content : bytes — raw PDF bytes

        Returns:
            str — concatenated text from all pages
        """
        if not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF required for PDF parsing: pip install PyMuPDF")

        text_parts = []

        # stream=content tells fitz to read from bytes, not a file path
        # filetype="pdf" is required when using stream mode
        with fitz.open(stream=content, filetype="pdf") as pdf_doc:
            for page_num, page in enumerate(pdf_doc):
                # get_text("text") extracts plain text (vs "html", "dict", "blocks")
                page_text = page.get_text("text")
                if page_text.strip():  # Skip blank pages
                    text_parts.append(page_text)

        return "\n".join(text_parts)

    def _parse_docx(self, content):
        """
        Extracts text from DOCX bytes using python-docx.

        How DOCX files work internally:
        - A .docx is actually a ZIP file containing XML files
        - python-docx unzips it and parses the XML
        - Text lives in <w:p> (paragraph) and <w:tc> (table cell) XML elements

        We extract from BOTH paragraphs AND table cells because resumes
        often use tables for formatting — missing table cells = missing skills!

        Args:
            content : bytes — raw DOCX bytes

        Returns:
            str — extracted text with newlines between paragraphs
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required: pip install python-docx")

        # BytesIO wraps bytes into a file-like object that Document() can open
        # Without BytesIO, Document() would expect a file path string
        doc = Document(io.BytesIO(content))
        text_parts = []

        # Extract paragraph text (main body content)
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract table cell text (skills/experience sometimes formatted as tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return "\n".join(text_parts)

    # ──────────────────────────────────────────────────────────────────────────
    # PRIVATE METHODS — text cleaning
    # ──────────────────────────────────────────────────────────────────────────

    def _clean_text(self, text):
        """
        Normalises raw extracted text for downstream NLP processing.

        Why clean? Raw PDF/DOCX text often contains:
        - Multiple consecutive spaces/tabs from layout formatting
        - Strange unicode characters (bullet symbols, em-dashes, etc.)
        - Inconsistent line breaks (3 newlines where 1 would do)
        - Null bytes and other binary artifacts

        Args:
            text : str — raw extracted text

        Returns:
            str — normalised, clean text
        """
        if not text:
            return ""

        # Replace Windows line endings (\r\n) with Unix (\n)
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Replace unicode bullet points and special dashes with plain text equivalents
        text = text.replace("\u2022", "•")   # bullet
        text = text.replace("\u2013", "-")   # en-dash
        text = text.replace("\u2014", "-")   # em-dash
        text = text.replace("\u2018", "'")   # left single quote
        text = text.replace("\u2019", "'")   # right single quote
        text = text.replace("\u201c", '"')   # left double quote
        text = text.replace("\u201d", '"')   # right double quote

        # Remove null bytes and other control characters (except newlines and tabs)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

        # Collapse multiple spaces/tabs into a single space
        text = re.sub(r"[ \t]+", " ", text)

        # Collapse more than 2 consecutive newlines into exactly 2
        # (preserves paragraph breaks without huge gaps)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)

        return text.strip()
